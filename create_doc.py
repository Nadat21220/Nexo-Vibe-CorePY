import os
import subprocess
import sys

def install_and_import():
    try:
        import docx
    except ImportError:
        print("Instalando python-docx...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx
    return docx

docx = install_and_import()
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Add Title
title = doc.add_heading('Resumen Detallado: NexoVibe Core', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Introducción
doc.add_heading('1. Descripción General del Proyecto', level=1)
doc.add_paragraph(
    "NexoVibe Core es la plataforma central y arquitectura backend/frontend para la gestión operativa de NexoVibe. "
    "El proyecto integra las interfaces de control para clientes y empleados, ofreciendo un dashboard de administración optimizado "
    "y una estructura modular. Su propósito es proveer herramientas de gestión de roles, calendario, socios y tareas."
)

# Equipo de Desarrollo
doc.add_heading('2. Equipo de Desarrollo', level=1)
doc.add_paragraph("Proyecto de Base de Datos y Arquitectura Core elaborado por:")
doc.add_paragraph("- Diego Castillo Mota", style='List Bullet')
doc.add_paragraph("- Diego Arellano", style='List Bullet')
doc.add_paragraph("- Jose Burgos", style='List Bullet')
doc.add_paragraph("- Anthony Poot", style='List Bullet')

# Tecnologías
doc.add_heading('3. Stack Tecnológico', level=1)
doc.add_heading('Frontend:', level=2)
doc.add_paragraph("- Framework: Next.js 14.2.35 (App Router)", style='List Bullet')
doc.add_paragraph("- Lenguaje: TypeScript", style='List Bullet')
doc.add_paragraph("- Estilos: Tailwind CSS, postcss", style='List Bullet')
doc.add_paragraph("- Dependencias Clave: React, Lucide-react (iconos), Recharts (gráficos), @hello-pangea/dnd (drag & drop para Kanban)", style='List Bullet')

doc.add_heading('Backend:', level=2)
doc.add_paragraph("- Lenguaje: Python", style='List Bullet')
doc.add_paragraph("- Framework: Flask 3.0.3, flask-cors", style='List Bullet')
doc.add_paragraph("- Driver DB: pg8000 1.31.5", style='List Bullet')
doc.add_paragraph("- Entorno: python-dotenv", style='List Bullet')
doc.add_paragraph("- Archivo principal: appi_nexovibe.py (API y conexión a BD)", style='List Bullet')

doc.add_heading('Base de Datos:', level=2)
doc.add_paragraph("- Motor: PostgreSQL", style='List Bullet')
doc.add_paragraph("- Estructura: Relacional (SQL)", style='List Bullet')
doc.add_paragraph("- Script de creación: nexovibe_bd.sql (contiene creación de base de datos 'nexovibe_bd', tablas y datos iniciales)", style='List Bullet')

# Características
doc.add_heading('4. Características Principales (Versión 1.0)', level=1)
doc.add_paragraph("- Dashboard Administrativo: Paneles de control separados para mejorar usabilidad.", style='List Bullet')
doc.add_paragraph("- Gestión de Roles: Interfaces y flujos específicos para Clientes y Empleados.", style='List Bullet')
doc.add_paragraph("- Autenticación (/api/auth): Gestión de acceso al sistema.", style='List Bullet')
doc.add_paragraph("- Calendario de actividades (/api/calendar): Gestión de eventos y programación.", style='List Bullet')
doc.add_paragraph("- Gestión de socios/clientes (/api/socios): Administración de la base de usuarios.", style='List Bullet')
doc.add_paragraph("- Administrador de tareas y Kanban (/api/tasks): Seguimiento visual y organizado del trabajo.", style='List Bullet')

# Arquitectura y Ejecución
doc.add_heading('5. Arquitectura, Instalación y Ejecución Local', level=1)
doc.add_paragraph("El proyecto se divide en tres capas principales que deben ejecutarse de manera coordinada:")
doc.add_heading('Capa de Base de Datos:', level=2)
doc.add_paragraph("1. Requiere PostgreSQL instalado.", style='List Number')
doc.add_paragraph("2. Configurar credenciales (usuario postgres / pwd default) o actualizar appi_nexovibe.py.", style='List Number')
doc.add_paragraph("3. Ejecutar nexovibe_bd.sql para inicializar la BD, tablas y datos semilla.", style='List Number')
doc.add_heading('Capa Backend (API Python):', level=2)
doc.add_paragraph("1. Instalar requerimientos: pip install -r requirements.txt", style='List Number')
doc.add_paragraph("2. Ejecutar: python appi_nexovibe.py", style='List Number')
doc.add_paragraph("3. Se despliega en http://localhost:5000", style='List Number')
doc.add_heading('Capa Frontend (Next.js):', level=2)
doc.add_paragraph("1. Instalar módulos: npm install", style='List Number')
doc.add_paragraph("2. Ejecutar servidor dev: npm run dev", style='List Number')
doc.add_paragraph("3. Se despliega en http://localhost:3000", style='List Number')

doc.add_heading('6. Estructura de Archivos Relevantes', level=1)
doc.add_paragraph("- /src: Contiene todo el código del Frontend Next.js (Componentes, App Router, Hooks).", style='List Bullet')
doc.add_paragraph("- appi_nexovibe.py: Servidor backend en Flask y lógica de la API.", style='List Bullet')
doc.add_paragraph("- nexovibe_bd.sql: Archivo de volcado y estructura SQL completo.", style='List Bullet')
doc.add_paragraph("- Read_preview.txt / README.md: Documentación y guías de ejecución locales.", style='List Bullet')
doc.add_paragraph("- package.json / requirements.txt: Listado de dependencias para JS y Python respectivamente.", style='List Bullet')

doc.add_heading('7. Estructura de la Base de Datos', level=1)
doc.add_paragraph("La base de datos PostgreSQL está modelada de manera relacional. A continuación se listan las tablas principales y de relación que conforman el sistema:")
doc.add_heading('Tablas Principales (Entidades):', level=2)
doc.add_paragraph("- persona: Entidad base con datos personales compartidos.", style='List Bullet')
doc.add_paragraph("- cliente: Datos específicos de los clientes, dependiente de 'persona'.", style='List Bullet')
doc.add_paragraph("- empleado: Datos específicos del personal de NexoVibe, dependiente de 'persona'.", style='List Bullet')
doc.add_paragraph("- creador_ugc: Datos de los creadores de contenido (User Generated Content).", style='List Bullet')
doc.add_paragraph("- servicio: Catálogo de servicios ofrecidos por la plataforma.", style='List Bullet')
doc.add_paragraph("- proyecto: Contiene la información principal, estado, prioridad y seguimiento de los proyectos.", style='List Bullet')
doc.add_paragraph("- tecnologia: Catálogo de tecnologías empleadas en los proyectos.", style='List Bullet')
doc.add_paragraph("- metrica: Definición de métricas para medir el rendimiento/KPIs.", style='List Bullet')

doc.add_heading('Tablas Intermedias (Relaciones N:M):', level=2)
doc.add_paragraph("- cliente_servicio: Relaciona qué clientes han contratado qué servicios.", style='List Bullet')
doc.add_paragraph("- cliente_proyecto: Vincula clientes con sus proyectos activos.", style='List Bullet')
doc.add_paragraph("- empleado_proyecto: Asignación de empleados a proyectos (gestión de equipo).", style='List Bullet')
doc.add_paragraph("- creadorugc_proyecto: Asignación de creadores UGC a proyectos de contenido.", style='List Bullet')
doc.add_paragraph("- proyecto_tecnologia: Tecnologías específicas utilizadas por cada proyecto.", style='List Bullet')
doc.add_paragraph("- proyecto_metrica: Métricas rastreadas en cada proyecto particular.", style='List Bullet')

doc.add_heading('8. Propósito de este documento', level=1)
doc.add_paragraph(
    "Este resumen está diseñado para proporcionar contexto a modelos de Inteligencia Artificial "
    "para que puedan comprender el alcance, las tecnologías y la arquitectura general del proyecto, "
    "permitiéndoles generar documentación adicional, manuales técnicos o asistir en el desarrollo."
)

import os
if not os.path.exists('documentacion'):
    os.makedirs('documentacion')

doc_path = os.path.join('documentacion', 'Resumen_Proyecto_NexoVibe.docx')
doc.save(doc_path)
print("Documento guardado en:", doc_path)
